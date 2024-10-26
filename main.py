import os
from docx import Document as DocxDocument
import win32com.client as win32
from PyPDF2 import PdfReader
from transformers import pipeline, AutoTokenizer, AutoModelForMaskedLM
from docx.shared import RGBColor
import re
from fuzzywuzzy import fuzz
import jieba
import jieba.posseg as pseg
from keybert import KeyBERT
from sklearn.feature_extraction.text import CountVectorizer

# 指定本地模型路径
model_path = r"C:\Users\15711\.cache\huggingface\hub\models--google-bert--bert-base-chinese\snapshots\c30a6ed22ab4564dc1e3b2ecbf6e766b0611a33f"

# 检查模型路径是否存在
if not os.path.exists(model_path):
    raise FileNotFoundError(f"Model path {model_path} does not exist.")

# 检查必要的文件是否存在
required_files = ["config.json", "pytorch_model.bin", "vocab.txt"]
for file in required_files:
    if not os.path.exists(os.path.join(model_path, file)):
        raise FileNotFoundError(f"Required file {file} not found in {model_path}.")

# 加载分词器和模型
tokenizer = AutoTokenizer.from_pretrained(model_path)
model = AutoModelForMaskedLM.from_pretrained(model_path)

# 创建中文填空预测模型
classifier = pipeline("fill-mask", model=model, tokenizer=tokenizer)

# 读取白名单
def read_whitelist(file_path):
    if not os.path.exists(file_path):
        return set()
    with open(file_path, 'r', encoding='utf-8') as f:
        return set(line.strip() for line in f)

# 写入白名单
def write_whitelist(file_path, whitelist):
    with open(file_path, 'w', encoding='utf-8') as f:
        for word in whitelist:
            f.write(f"{word}\n")

# 对一段文字中的每个字进行掩码并预测
# 对一段文字中的每个字进行掩码并预测
def mask_and_predict_all_words(text, whitelist, max_length=510):
    errors = []
    predictions = []

    # 将文本分成多个段落，每个段落不超过 max_length 个 token
    segments = [text[i:i + max_length] for i in range(0, len(text), max_length)]

    for segment in segments:
        for i in range(len(segment)):
            # 跳过标点符号和阿拉伯数字
            if re.match(r'\d+|[^\w\s]', segment[i]):
                continue

            # 检查是否是量词加后面的字
            if is_quantifier_word(segment, i):
                continue

            masked_text = segment[:i] + "[MASK]" + segment[i + 1:]
            original_word = segment[i]

            # 进行预测
            prediction = classifier(masked_text, top_k=2)

            # 检查预测结果是否包含原始字
            predicted_words = [pred['token_str'] for pred in prediction]
            if original_word not in predicted_words[:2]:  # 检查前两个预测结果是否包含原始字
                # 进行模糊搜索
                for word in whitelist:
                    if fuzz.partial_ratio(original_word, word) > 80:  # 设置模糊匹配的阈值
                        break
                else:
                    errors.append((i, original_word))

            predictions.append((i, original_word, predicted_words))

    return errors, predictions

# 检查是否是量词加后面的字
def is_quantifier_word(segment, index):
    # 检查是否是量词和汉字组合的词
    if re.match(r'[一二三四五六七八九十百千万亿两]+[\u4e00-\u9fa5]+', segment[index]):
        return True
    # 检查是否是量词加后面两个字
    if index + 2 < len(segment) and re.match(r'[一二三四五六七八九十百千万亿两]+[\u4e00-\u9fa5]{2}', segment[index:index + 3]):
        return True
    return False

# 检查是否是人名
def is_name(segment, index):
    # 使用 jieba 进行分词和词性标注
    words = pseg.lcut(segment)
    for word, flag in words:
        if flag == 'nr' and word == segment[index]:
            return True
    return False

# 读取 .docx 文件内容
def read_docx_content(file_path):
    # 打开文档
    doc = DocxDocument(file_path)

    # 初始化存储所有页内容的列表
    pages_content = []

    # 遍历文档中的所有段落
    for paragraph in doc.paragraphs:
        # 获取段落文本
        text = paragraph.text

        # 如果段落文本不为空，则添加到当前页的内容中
        if text.strip():
            pages_content.append(text)

    return pages_content

# 读取 .doc 文件内容
def read_doc_content(file_path):
    # 打开文档
    word = win32.Dispatch("Word.Application")
    word.Visible = False
    doc = word.Documents.Open(file_path)

    # 初始化存储所有页内容的列表
    pages_content = []

    # 遍历文档中的所有段落
    for paragraph in doc.Paragraphs:
        # 获取段落文本
        text = paragraph.Range.Text

        # 如果段落文本不为空，则添加到当前页的内容中
        if text.strip():
            pages_content.append(text)

    # 关闭文档
    doc.Close()
    word.Quit()

    return pages_content

# 读取 .pdf 文件内容
def read_pdf_content(file_path):
    # 打开文档
    pdf_file = open(file_path, 'rb')
    pdf_reader = PdfReader(pdf_file)

    # 初始化存储所有页内容的列表
    pages_content = []

    # 遍历文档中的所有页
    for page_num in range(len(pdf_reader.pages)):
        page = pdf_reader.pages[page_num]
        text = page.extract_text()

        # 如果页文本不为空，则添加到当前页的内容中
        if text.strip():
            pages_content.append(text)

    # 关闭文档
    pdf_file.close()

    return pages_content

# 处理文档
def process_document_file(file_path, permanent_whitelist_path, temporary_whitelist_path, file_type, compare_temporary_whitelist):
    # 读取文档内容
    if file_type == "docx":
        pages_content = read_docx_content(file_path)
    elif file_type == "doc":
        pages_content = read_doc_content(file_path)
    elif file_type == "pdf":
        pages_content = read_pdf_content(file_path)
    else:
        raise ValueError(f"不支持的文件类型: {file_type}")

    # 读取永久白名单和临时白名单
    permanent_whitelist = read_whitelist(permanent_whitelist_path)
    temporary_whitelist = read_whitelist(temporary_whitelist_path)

    # 合并白名单
    if compare_temporary_whitelist:
        combined_whitelist = permanent_whitelist.union(temporary_whitelist)
    else:
        combined_whitelist = permanent_whitelist

    # 处理每一页的内容
    all_errors = []
    for i, page in enumerate(pages_content):
        print(f"Page {i + 1}:")
        print(page)
        print()

        # 对每一页的内容进行全部掩码和预测
        errors, predictions = mask_and_predict_all_words(page, combined_whitelist)
        all_errors.extend(errors)

        # 打印错误和预测结果
        print(f"Errors in Page {i + 1}:")
        for error in errors:
            print(f"Index: {error[0]}, Original Word: {error[1]}")
        print()

        print(f"Predictions in Page {i + 1}:")
        for pred in predictions:
            print(f"Index: {pred[0]}, Original Word: {pred[1]}, Predicted Words: {pred[2]}")
        print()

    # 将预测错误的字标红并写回文档中
    write_errors_to_docx(pages_content, all_errors, file_path, file_type)

    # 询问是否保存临时白名单到永久白名单
    save_temporary_whitelist = input("是否将临时白名单保存到永久白名单？(y/n): ").strip().lower()
    if save_temporary_whitelist == 'y':
        permanent_whitelist.update(temporary_whitelist)
        write_whitelist(permanent_whitelist_path, permanent_whitelist)
        print(f"临时白名单已保存到永久白名单: {permanent_whitelist_path}")
    else:
        # 清空临时白名单
        write_whitelist(temporary_whitelist_path, set())
        print("临时白名单已清空。")

# 将预测错误的字标红并写回文档中
def write_errors_to_docx(pages_content, all_errors, file_path, file_type):
    # 创建一个新的 Word 文档
    doc = DocxDocument()

    # 逐段处理文本，将正确的字和错误的字分别写入文档中
    for page in pages_content:
        paragraph = doc.add_paragraph()
        for i, char in enumerate(page):
            if (i, char) in all_errors:
                # 将错误的字标红
                run = paragraph.add_run(char)
                run.font.color.rgb = RGBColor(255, 0, 0)  # 红色
            else:
                # 将正确的字写入
                run = paragraph.add_run(char)

    # 保存文档
    output_file_path = file_path.replace(".doc", "_errors.doc").replace(".docx", "_errors.docx").replace(".pdf", "_errors.docx")
    doc.save(output_file_path)
    print(f"预测错误的字已标红并保存到 {output_file_path}")

def main():
    original_file_path = None
    original_permanent_whitelist_path = None
    original_temporary_whitelist_path = None
    original_file_type = None

    while True:
        # 提示用户输入文件类型
        if original_file_type is None:
            file_type = input("请输入文件类型 (doc/docx/pdf): ").strip().lower()
            if file_type not in ["doc", "docx", "pdf"]:
                print("不支持的文件类型，请输入 'doc'、'docx' 或 'pdf'。")
                continue
        else:
            file_type = original_file_type

        # 提示用户输入文件的完整路径
        if original_file_path is None:
            file_path = input("请输入文件的完整路径(无双引号): ").strip()
        else:
            file_path = original_file_path

        # 提示用户输入永久白名单文件的完整路径
        if original_permanent_whitelist_path is None:
            permanent_whitelist_path = input("请输入永久白名单文件的完整路径(无双引号): ").strip()
        else:
            permanent_whitelist_path = original_permanent_whitelist_path

        # 提示用户输入临时白名单文件的完整路径
        if original_temporary_whitelist_path is None:
            temporary_whitelist_path = input("请输入临时白名单文件的完整路径(无双引号): ").strip()
        else:
            temporary_whitelist_path = original_temporary_whitelist_path

        # 检查文件是否存在
        if not os.path.exists(file_path):
            print(f"文件 {file_path} 不存在。")
            continue

        # 检查永久白名单文件是否存在
        if not os.path.exists(permanent_whitelist_path):
            print(f"永久白名单文件 {permanent_whitelist_path} 不存在。")
            continue

        # 检查文件类型和文件路径是否匹配
        file_extension = os.path.splitext(file_path)[1].lower()
        if file_extension != f".{file_type}":
            print(f"文件类型和文件路径不匹配。选择的文件类型是 {file_type}，但文件路径的扩展名是 {file_extension}。")
            continue

        # 读取文档内容
        if file_type == "docx":
            pages_content = read_docx_content(file_path)
        elif file_type == "doc":
            pages_content = read_doc_content(file_path)
        elif file_type == "pdf":
            pages_content = read_pdf_content(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {file_type}")

        # 提取关键词
        text = '\n'.join(pages_content)
        keywords = extract_keywords(text)
        print(f"提取的关键词: {keywords}")

        # 将提取的关键词写入临时白名单
        temporary_whitelist = read_whitelist(temporary_whitelist_path)
        temporary_whitelist.update(keywords)
        write_whitelist(temporary_whitelist_path, temporary_whitelist)
        print(f"关键词已写入临时白名单: {temporary_whitelist_path}")

        # 询问是否要对临时白名单进行比较
        compare_temporary_whitelist = input("是否要对临时白名单进行比较？(y/n): ").strip().lower()
        if compare_temporary_whitelist == 'y':
            compare_temporary_whitelist = True
        else:
            compare_temporary_whitelist = False

        # 处理文档
        process_document_file(file_path, permanent_whitelist_path, temporary_whitelist_path, file_type, compare_temporary_whitelist)

        # 程序结束时询问是否要向白名单添加特色词
        while True:
            add_special_words = input("程序结束，是否要向白名单添加特色词？(y/n/q): ").strip().lower()
            if add_special_words == 'q':
                break
            elif add_special_words == 'y':
                permanent_whitelist = read_whitelist(permanent_whitelist_path)
                while True:
                    special_word = input("请输入要添加的特色词(输入'q'退出): ").strip()
                    if special_word == 'q':
                        break
                    permanent_whitelist.add(special_word)
                write_whitelist(permanent_whitelist_path, permanent_whitelist)
            elif add_special_words == 'n':
                break

        # 询问是否重新进行查错
        restart = input("是否要重新进行查错？(y/n): ").strip().lower()
        if restart != 'y':
            break

        # 询问是否是原来的文件
        same_file = input("是否是原来的文件？(y/n): ").strip().lower()
        if same_file == 'y':
            original_file_path = file_path
            original_permanent_whitelist_path = permanent_whitelist_path
            original_temporary_whitelist_path = temporary_whitelist_path
            original_file_type = file_type
        else:
            original_file_path = None
            original_permanent_whitelist_path = None
            original_temporary_whitelist_path = None
            original_file_type = None

def extract_keywords(text):
    # 本地模型路径
    model_path = r"C:\Users\15711\.cache\huggingface\hub\models--sentence-transformers--all-MiniLM-L6-v2\snapshots\8b3219a92973c328a8e22fadcfa821b5dc75636a"

    # 初始化KeyBERT模型
    kw_model = KeyBERT(model=model_path)

    # 使用jieba进行中文分词
    def chinese_tokenizer(text):
        return list(jieba.cut(text))

    # 使用CountVectorizer进行分词
    vectorizer = CountVectorizer(tokenizer=chinese_tokenizer, ngram_range=(1, 1))

    # 提取关键词
    keywords = kw_model.extract_keywords(text, vectorizer=vectorizer, keyphrase_ngram_range=(1, 1), top_n=30)

    # 返回关键词列表
    return [keyword[0] for keyword in keywords]

if __name__ == "__main__":
    main()






























